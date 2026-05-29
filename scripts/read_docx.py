#!/usr/bin/env python3
"""Extract text + tables from a DOCX file.

Usage:
    read_docx.py file.docx               # full text + tables as markdown
    read_docx.py file.docx --json        # structured JSON
    read_docx.py file.docx --no-tables   # skip tables
"""
import argparse
import json
import sys


def docx_to_md(path, include_tables=True):
    from docx import Document
    doc = Document(path)
    blocks = []

    # Preserve order of paragraphs and tables
    body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)
    paras_by_xml = {p._element: p for p in doc.paragraphs}
    tables_by_xml = {t._element: t for t in doc.tables}

    for el in body.iterchildren():
        if el in paras_by_xml:
            p = paras_by_xml[el]
            text = p.text.strip()
            if not text:
                continue
            style = (p.style.name or "").lower() if p.style else ""
            if style.startswith("heading 1") or style == "title":
                blocks.append({"type": "heading", "level": 1, "text": text})
            elif style.startswith("heading 2"):
                blocks.append({"type": "heading", "level": 2, "text": text})
            elif style.startswith("heading"):
                blocks.append({"type": "heading", "level": 3, "text": text})
            else:
                blocks.append({"type": "paragraph", "text": text})
        elif include_tables and el in tables_by_xml:
            t = tables_by_xml[el]
            rows = [[c.text.strip() for c in row.cells] for row in t.rows]
            blocks.append({"type": "table", "rows": rows})
    return blocks


def render_md(blocks):
    out = []
    for b in blocks:
        if b["type"] == "heading":
            out.append("#" * b["level"] + " " + b["text"])
        elif b["type"] == "paragraph":
            out.append(b["text"])
        elif b["type"] == "table":
            rows = b["rows"]
            if not rows:
                continue
            head = rows[0]
            out.append("| " + " | ".join(head) + " |")
            out.append("|" + "|".join("---" for _ in head) + "|")
            for r in rows[1:]:
                out.append("| " + " | ".join(c.replace("\n", " ") for c in r) + " |")
        out.append("")
    return "\n".join(out)


def main():
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("docx")
    ap.add_argument("--json", action="store_true")
    ap.add_argument("--no-tables", action="store_true")
    a = ap.parse_args()
    blocks = docx_to_md(a.docx, include_tables=not a.no_tables)
    if a.json:
        print(json.dumps(blocks, indent=2))
    else:
        print(render_md(blocks))


if __name__ == "__main__":
    try:
        main()
    except FileNotFoundError as e:
        print(f"ERROR: {e}", file=sys.stderr); sys.exit(2)
